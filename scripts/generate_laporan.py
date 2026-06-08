"""
Generator Laporan Word EDUFIN
Membuat laporan dalam format template UB berdasarkan data project
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os
import re

# ─────────────────────────────────────────
# DATA PROJECT EDUFIN
# ─────────────────────────────────────────
PROJECT_DATA = {
    "judul": "EDUFIN – Platform Manajemen Keuangan Pendidikan",
    "semester": "LAPORAN PROJECT UTS SEMESTER GENAP 2025/2026",
    "kelompok": "Kelompok 5 / Kelas D3TI-B",
    "anggota": [
        {"nim": "25314071111062", "nama": "Diego Armando Ramadan", "jobdesc": "Frontend Development (80%) – Implementasi UI/UX, komponen React, halaman siswa, donatur, dan admin sekolah menggunakan Tailwind CSS v4. Integrasi LocalStorage mock data."},
        {"nim": "253140707111017", "nama": "David Bimantoro Sarashadi", "jobdesc": "Backend Development (85%) – Desain arsitektur database Supabase (PostgreSQL), konfigurasi RLS, skema tabel, Edge Functions (Xendit webhook, WhatsApp notifikasi), seed data."},
    ],
    "prodi": "PROGRAM STUDI D3 TEKNOLOGI INFORMASI",
    "fakultas": "FAKULTAS VOKASI",
    "universitas": "UNIVERSITAS BRAWIJAYA",
    "kota": "MALANG",
    "tahun": "2026",
    "deskripsi": (
        "EDUFIN adalah platform manajemen keuangan pendidikan berbasis web yang dirancang untuk "
        "mendigitalisasi proses pembayaran SPP dan penggalangan dana di lingkungan sekolah. "
        "Platform ini menghubungkan tiga pihak utama: Admin Sekolah yang mengelola tagihan dan data siswa, "
        "Siswa/Orang Tua yang melakukan pembayaran SPP secara digital, serta Donatur yang dapat berkontribusi "
        "pada kampanye penggalangan dana siswa.\n\n"
        "Teknologi yang digunakan: React 18 + TypeScript, Vite 6, Tailwind CSS v4, Supabase (PostgreSQL + Auth + Edge Functions), "
        "dan Xendit sebagai Payment Gateway untuk QRIS, Virtual Account, dan E-Wallet. "
        "Platform ini juga terintegrasi dengan WhatsApp (Fonnte/Wablas) untuk notifikasi otomatis "
        "pengingat tagihan dan konfirmasi pembayaran kepada orang tua siswa."
    ),
}

# ─────────────────────────────────────────
# DATA FITUR & MODUL
# ─────────────────────────────────────────
FITUR_MODUL = [
    {
        "label": "Auth & Onboarding",
        "items": [
            ("Login", "Halaman masuk dengan email & password, validasi format, checkbox 'Ingat Saya'. Mendukung multi-role redirect (siswa→PWA, sekolah→desktop, donatur→web)."),
            ("Register", "Pendaftaran akun untuk Donatur dan Siswa/Orang Tua. Siswa harus mencocokkan NISN dan email orang tua yang diimpor sekolah."),
            ("Beranda / Onboarding", "Splash screen yang menampilkan pilihan role sebelum diarahkan ke halaman login yang sesuai."),
            ("Lupa & Reset Password", "Pengiriman tautan reset sandi ke email. Validasi kekuatan sandi baru (min. 8 karakter, 1 huruf kapital, 1 angka)."),
        ]
    },
    {
        "label": "User – Siswa / Orang Tua (Mobile PWA)",
        "items": [
            ("Dashboard Siswa", "Ringkasan tagihan aktif, saldo tabungan, dan tombol cepat bayar SPP. Terhubung ke notifikasi in-app."),
            ("Bayar SPP", "Daftar tagihan (Belum Bayar, Terlambat, Lunas, Cicilan). Pilih metode pembayaran: QRIS, VA, E-Wallet, atau transfer manual."),
            ("Riwayat Pembayaran", "Histori transaksi SPP dan unduhan tanda terima (receipt) PDF."),
            ("Pengajuan Cicilan", "Form pengajuan cicilan SPP (2x–6x), pemantauan status persetujuan, dan pembayaran per periode."),
            ("Kampanye Penggalangan Dana", "Buat dan pantau kampanye donasi pribadi beserta progres dana terkumpul."),
            ("Profil Siswa", "Kartu identitas siswa (NISN, nama, kelas), edit kontak orang tua, dan ganti password."),
        ]
    },
    {
        "label": "Admin Sekolah (Desktop Dashboard)",
        "items": [
            ("Dashboard Keuangan", "Statistik ringkasan: total siswa, penerimaan SPP bulan ini, tunggakan, kampanye pending. Grafik bar bulanan."),
            ("Manajemen Tagihan (CRUD)", "Buat tagihan SPP massal per kelas/individu, edit, batalkan, verifikasi transfer manual, input pembayaran tunai."),
            ("Manajemen Siswa (CRUD)", "Daftar siswa, tambah/edit/nonaktifkan, impor massal dari CSV dengan pratinjau validasi."),
            ("Manajemen Beasiswa (CRUD)", "Buat program beasiswa, alokasikan penerima dari siswa aktif, pantau kuota dan status beasiswa."),
            ("Manajemen Kampanye", "Tinjau, setujui/tolak kampanye siswa, dan minta pencairan dana kampanye yang telah selesai."),
            ("Manajemen Cicilan", "Tinjau persetujuan cicilan dari siswa, setujui atau tolak dengan alasan."),
            ("Laporan Keuangan", "Grafik SPP vs Donasi per bulan, rincian tunggakan per kelas, ekspor PDF laporan (V2)."),
            ("WhatsApp Blast", "Kirim pesan massal ke orang tua siswa seluruh sekolah atau kelas tertentu."),
            ("Profil Sekolah", "Edit logo, rekening bank pencairan donasi, dan preferensi notifikasi sekolah."),
        ]
    },
    {
        "label": "Donatur (Web/Mobile)",
        "items": [
            ("Eksplorasi Kampanye", "Jelajahi kampanye aktif dari seluruh sekolah, filter berdasarkan kategori/sekolah."),
            ("Detail Kampanye & Donasi", "Lihat deskripsi, progres dana, daftar donatur. Donasi mulai Rp 10.000 (anonim/publik) via Xendit."),
            ("Transparansi Feed (V2)", "Lihat laporan penggunaan dana yang diunggah oleh siswa penerima kampanye."),
            ("Profil Donatur", "Edit nama publik, email, telepon. Ringkasan total dan histori donasi."),
        ]
    },
    {
        "label": "Super Admin EDUFIN (Panel Internal)",
        "items": [
            ("Platform Overview", "Statistik nasional: total sekolah mitra, total siswa, GMV transaksi, kampanye aktif. Grafik pertumbuhan bulanan."),
            ("Manajemen Sekolah", "Onboarding sekolah baru, edit profil sekolah, suspend/aktifkan kembali tenant."),
            ("Moderasi Kampanye", "Pantau dan bekukan kampanye terindikasi fraud dari seluruh sekolah. Log audit moderasi."),
            ("Impersonasi Admin", "Masuk sebagai Admin Sekolah tertentu untuk keperluan dukungan teknis."),
        ]
    },
]

# ─────────────────────────────────────────
# DATA PERANCANGAN PER MODUL
# ─────────────────────────────────────────
PERANCANGAN_MODUL = [
    {
        "nama": "Autentikasi & Manajemen Pengguna",
        "identifikasi_user": [
            ("Super Admin (EDUFIN)", "Super Admin adalah pihak internal EDUFIN yang mendaftarkan sekolah baru dan membuat akun admin sekolah pertama ke dalam sistem."),
            ("Admin Sekolah", "Admin Sekolah adalah pengguna yang mendapatkan akun dari Super Admin dan dapat mengelola data siswa serta mengundang orang tua bergabung."),
            ("Siswa / Orang Tua", "Siswa/Orang Tua adalah pengguna aplikasi mobile PWA untuk membayar SPP dan mengelola kampanye. Akun dibuat melalui tautan undangan dari sekolah."),
            ("Donatur", "Donatur adalah pengguna publik yang dapat mendaftar secara mandiri menggunakan email atau Google OAuth dan melakukan donasi ke kampanye siswa."),
        ],
        "use_case": [
            ("Pengguna harus dapat masuk (login) menggunakan email dan kata sandi.", "Super Admin, Admin Sekolah, Siswa/Orang Tua, Donatur"),
            ("Pengguna harus dapat keluar (logout) dari sesi aktif.", "Semua Role"),
            ("Sistem harus membedakan akses halaman berdasarkan role pengguna (Protected Route).", "Semua Role"),
            ("Admin Sekolah dapat mendaftarkan diri setelah menerima undangan dari Super Admin.", "Admin Sekolah"),
            ("Siswa/Orang Tua dapat mendaftar setelah menerima tautan undangan yang dikirim sekolah.", "Siswa / Orang Tua"),
            ("Donatur dapat mendaftar secara mandiri menggunakan email atau Google OAuth.", "Donatur"),
            ("Pengguna dapat meminta tautan reset kata sandi melalui email (Lupa Password).", "Semua Role"),
            ("Pengguna dapat mengganti kata sandi dari halaman pengaturan profil.", "Semua Role"),
            ("Sistem harus memverifikasi kekuatan kata sandi baru (min. 8 karakter, 1 huruf kapital, 1 angka).", "Semua Role"),
            ("Admin Sekolah dapat menambahkan sub-admin dengan izin granular (Full Admin, Finance Only, View Only).", "Admin Sekolah"),
        ],
        "normalisasi": [
            ("auth.users (Supabase Built-in)", [
                ("id", "UUID", "Primary Key"),
                ("email", "VARCHAR", "Email unik pengguna"),
                ("encrypted_password", "TEXT", "Sandi terenkripsi (bcrypt)"),
                ("created_at", "TIMESTAMP", "Waktu akun dibuat"),
            ]),
            ("school_admins", [
                ("id", "UUID", "Primary Key"),
                ("user_id", "UUID", "FK → auth.users.id"),
                ("school_id", "UUID", "FK → schools.id"),
                ("name", "VARCHAR", "Nama lengkap admin"),
                ("role", "VARCHAR", "Peran: 'Full Admin', 'Finance Only', 'View Only'"),
                ("permissions", "JSONB", "Hak akses granular"),
                ("is_super_admin", "BOOLEAN", "TRUE jika Super Admin EDUFIN"),
            ]),
        ],
        "relasi": [
            ("school_admins", "user_id", "auth.users", "id", "Many-to-One"),
            ("school_admins", "school_id", "schools", "id", "Many-to-One"),
            ("students", "user_id", "auth.users", "id", "One-to-One (nullable)"),
        ],
    },
    {
        "nama": "Manajemen Pembayaran SPP",
        "identifikasi_user": [
            ("Admin Sekolah", "Admin Sekolah adalah pengguna yang membuat dan mengelola tagihan SPP, memverifikasi bukti transfer manual, dan mencatat pembayaran tunai."),
            ("Siswa / Orang Tua", "Siswa/Orang Tua adalah pengguna yang melihat tagihan SPP aktif, memilih metode pembayaran, melakukan pembayaran, dan mengunduh bukti pembayaran."),
        ],
        "use_case": [
            ("Admin Sekolah dapat membuat tagihan SPP untuk satu siswa atau seluruh kelas secara massal.", "Admin Sekolah"),
            ("Admin Sekolah dapat menetapkan nominal dan tanggal jatuh tempo pada setiap tagihan.", "Admin Sekolah"),
            ("Admin Sekolah dapat menyetujui atau menolak bukti transfer manual yang diunggah siswa.", "Admin Sekolah"),
            ("Admin Sekolah dapat menginput pembayaran tunai (cash) secara manual.", "Admin Sekolah"),
            ("Siswa/Orang Tua dapat melihat daftar tagihan SPP beserta statusnya (Lunas, Belum Bayar, Terlambat, Cicilan).", "Siswa / Orang Tua"),
            ("Siswa/Orang Tua dapat memilih tagihan dan melakukan pembayaran via QRIS, Virtual Account, atau E-Wallet.", "Siswa / Orang Tua"),
            ("Siswa/Orang Tua dapat mengunggah foto bukti transfer untuk pembayaran manual.", "Siswa / Orang Tua"),
            ("Siswa/Orang Tua dapat mengunduh tanda terima (receipt) pembayaran dalam format PDF.", "Siswa / Orang Tua"),
            ("Sistem harus memperbarui status tagihan secara otomatis setelah pembayaran berhasil diterima via Xendit.", "Sistem"),
        ],
        "normalisasi": [
            ("bills", [
                ("id", "UUID", "Primary Key"),
                ("school_id", "UUID", "FK → schools.id"),
                ("student_id", "UUID", "FK → students.id"),
                ("amount", "INTEGER", "Nominal tagihan SPP"),
                ("late_fee", "INTEGER", "Nominal denda keterlambatan"),
                ("month", "VARCHAR", "Periode tagihan (e.g., 'Juni 2026')"),
                ("due_date", "DATE", "Tanggal jatuh tempo"),
                ("status", "ENUM", "Status: 'lunas', 'belum_bayar', 'terlambat', 'cicilan'"),
                ("payment_method", "ENUM", "Metode: 'qris', 'va_bca', 'gopay', 'transfer', 'tunai'"),
                ("xendit_invoice_id", "VARCHAR", "ID invoice dari Xendit (nullable)"),
            ]),
        ],
        "relasi": [
            ("bills", "school_id", "schools", "id", "Many-to-One"),
            ("bills", "student_id", "students", "id", "Many-to-One"),
            ("installments", "bill_id", "bills", "id", "One-to-One"),
        ],
    },
    {
        "nama": "Manajemen Cicilan",
        "identifikasi_user": [
            ("Siswa / Orang Tua", "Siswa/Orang Tua adalah pengguna yang mengajukan permohonan cicilan SPP, memilih tenor, menyertakan alasan, dan membayar cicilan per periode."),
            ("Admin Sekolah", "Admin Sekolah adalah pengguna yang meninjau pengajuan cicilan, memberikan persetujuan atau penolakan, serta memantau kemajuan pembayaran cicilan aktif."),
        ],
        "use_case": [
            ("Siswa/Orang Tua dapat mengajukan cicilan untuk tagihan SPP yang belum lunas.", "Siswa / Orang Tua"),
            ("Siswa/Orang Tua dapat memilih jumlah periode cicilan (2x, 3x, 4x, 5x, atau 6x).", "Siswa / Orang Tua"),
            ("Siswa/Orang Tua dapat menyertakan alasan permohonan cicilan secara tertulis.", "Siswa / Orang Tua"),
            ("Siswa/Orang Tua dapat melihat rincian sub-tagihan per periode beserta tanggal jatuh tempo.", "Siswa / Orang Tua"),
            ("Admin Sekolah dapat melihat seluruh pengajuan cicilan yang masih berstatus menunggu (Pending).", "Admin Sekolah"),
            ("Admin Sekolah dapat menyetujui cicilan sehingga tagihan dipecah otomatis ke X periode.", "Admin Sekolah"),
            ("Admin Sekolah dapat menolak cicilan dengan mencantumkan alasan penolakan.", "Admin Sekolah"),
            ("Sistem harus mengubah status cicilan menjadi 'Selesai' saat semua periode telah lunas.", "Sistem"),
        ],
        "normalisasi": [
            ("installments", [
                ("id", "UUID", "Primary Key"),
                ("bill_id", "UUID", "FK → bills.id (tagihan asli)"),
                ("student_id", "UUID", "FK → students.id"),
                ("total_periods", "INTEGER", "Total periode cicilan (2–6)"),
                ("amount_per_period", "INTEGER", "Nominal per periode"),
                ("reason", "TEXT", "Alasan pengajuan cicilan"),
                ("status", "ENUM", "Status: 'pending_approval', 'active', 'completed', 'rejected'"),
            ]),
            ("installment_periods", [
                ("id", "UUID", "Primary Key"),
                ("installment_id", "UUID", "FK → installments.id"),
                ("period_number", "INTEGER", "Urutan periode (1, 2, 3, …)"),
                ("amount", "INTEGER", "Nominal cicilan periode ini"),
                ("due_date", "DATE", "Tanggal jatuh tempo periode"),
                ("status", "ENUM", "Status: 'belum_bayar', 'lunas', 'terlambat'"),
            ]),
        ],
        "relasi": [
            ("installments", "bill_id", "bills", "id", "One-to-One"),
            ("installments", "student_id", "students", "id", "Many-to-One"),
            ("installment_periods", "installment_id", "installments", "id", "Many-to-One"),
        ],
    },
    {
        "nama": "Manajemen Kampanye Penggalangan Dana",
        "identifikasi_user": [
            ("Siswa / Orang Tua", "Siswa/Orang Tua adalah pengguna yang membuat kampanye, mengunggah dokumen pendukung, memantau kemajuan donasi, dan melaporkan penggunaan dana."),
            ("Admin Sekolah", "Admin Sekolah adalah penjaga gerbang (gatekeeper) yang meninjau, menyetujui/menolak pengajuan kampanye dan meminta pencairan dana."),
            ("Donatur", "Donatur adalah pengguna publik yang menjelajahi kampanye aktif dan melakukan donasi secara digital (dengan pilihan anonim)."),
            ("Super Admin (EDUFIN)", "Super Admin memiliki wewenang untuk membekukan kampanye yang terindikasi kecurangan dari seluruh sekolah di platform."),
        ],
        "use_case": [
            ("Siswa/Orang Tua dapat membuat kampanye dengan mengisi judul, deskripsi, target nominal, durasi, dan kategori.", "Siswa / Orang Tua"),
            ("Admin Sekolah dapat menyetujui kampanye agar tayang secara publik.", "Admin Sekolah"),
            ("Admin Sekolah dapat meminta pencairan dana saat kampanye selesai/mencapai target.", "Admin Sekolah"),
            ("Donatur dapat menjelajahi dan mencari kampanye aktif dari seluruh sekolah.", "Donatur"),
            ("Donatur dapat melakukan donasi minimal Rp 10.000 dengan pilihan anonim.", "Donatur"),
            ("Sistem harus menutup kampanye secara otomatis saat batas waktu berakhir.", "Sistem"),
            ("Super Admin dapat membekukan (suspend) kampanye yang terindikasi fraud.", "Super Admin"),
        ],
        "normalisasi": [
            ("campaigns", [
                ("id", "UUID", "Primary Key"),
                ("school_id", "UUID", "FK → schools.id"),
                ("student_id", "UUID", "FK → students.id"),
                ("title", "VARCHAR", "Judul kampanye"),
                ("target_amount", "INTEGER", "Target dana yang ingin dikumpulkan"),
                ("current_amount", "INTEGER", "Total donasi yang sudah terkumpul"),
                ("status", "ENUM", "Status: 'pending', 'approved', 'rejected', 'completed', 'expired', 'suspended'"),
                ("end_date", "DATE", "Tanggal kampanye berakhir"),
            ]),
            ("donations", [
                ("id", "UUID", "Primary Key"),
                ("campaign_id", "UUID", "FK → campaigns.id"),
                ("donor_user_id", "UUID", "FK → auth.users.id (nullable untuk guest)"),
                ("amount", "INTEGER", "Jumlah donasi"),
                ("is_anonymous", "BOOLEAN", "TRUE jika donatur memilih anonim"),
                ("payment_status", "ENUM", "Status: 'pending', 'success', 'failed', 'expired'"),
            ]),
        ],
        "relasi": [
            ("campaigns", "school_id", "schools", "id", "Many-to-One"),
            ("campaigns", "student_id", "students", "id", "Many-to-One"),
            ("donations", "campaign_id", "campaigns", "id", "Many-to-One"),
            ("donations", "donor_user_id", "auth.users", "id", "Many-to-One (nullable)"),
        ],
    },
]


# ─────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_para(doc, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None):
    para = doc.add_paragraph()
    para.alignment = align
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    run = para.add_run(text)
    set_font(run, bold=bold, size=size)
    return para

def set_table_style(table):
    """Terapkan border tabel dan font Times New Roman"""
    table.style = 'Table Grid'

def add_table_row(table, values, bold=False, is_header=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_header or i == 0) else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        set_font(run, bold=bold or is_header, size=11)
        # Warna header
        if is_header:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'BDD7EE')  # Biru muda
            tcPr.append(shd)

def add_section_break(doc):
    from docx.enum.section import WD_SECTION
    doc.add_section(WD_SECTION.NEW_PAGE)

def add_heading_numbered(doc, text, bold=True, size=12):
    """Tambah heading dengan format bold 12pt TNR"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, bold=bold, size=size)
    return para

def add_placeholder_image(doc, caption):
    """Tambah placeholder untuk diagram/gambar"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {caption} ]")
    set_font(run, bold=False, size=11)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    # Border tabel placeholder
    return p


# ─────────────────────────────────────────
# GENERATE DOCUMENT
# ─────────────────────────────────────────
def generate_laporan():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)

    # ── HALAMAN COVER ──────────────────────────────────
    doc.add_paragraph()
    add_para(doc, PROJECT_DATA["judul"], bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, PROJECT_DATA["semester"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Oleh :", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, PROJECT_DATA["kelompok"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for a in PROJECT_DATA["anggota"]:
        add_para(doc, f"{a['nim']} – {a['nama']}", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, PROJECT_DATA["prodi"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, PROJECT_DATA["fakultas"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, PROJECT_DATA["universitas"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, PROJECT_DATA["kota"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, PROJECT_DATA["tahun"], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── HALAMAN BARU: JOBDESC ──────────────────────────
    doc.add_page_break()
    add_heading_numbered(doc, "Judul Project")
    add_para(doc, PROJECT_DATA["judul"], bold=False)
    doc.add_paragraph()

    add_heading_numbered(doc, "Jobdesc / Kinerja Anggota Kelompok")
    doc.add_paragraph()

    table_jobdesc = doc.add_table(rows=1, cols=4)
    set_table_style(table_jobdesc)
    add_table_row(table_jobdesc, ["NO", "NIM", "NAMA", "JOB DESC / KINERJA"], is_header=True)
    for i, a in enumerate(PROJECT_DATA["anggota"], 1):
        add_table_row(table_jobdesc, [str(i), a["nim"], a["nama"], a["jobdesc"]])

    # Atur lebar kolom
    for row in table_jobdesc.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(4.0)
        row.cells[3].width = Cm(8.0)

    doc.add_paragraph()

    # ── DESKRIPSI PROJECT ──────────────────────────────
    add_heading_numbered(doc, "Deskripsi Project / Overview")
    add_para(doc, PROJECT_DATA["deskripsi"], bold=False)
    doc.add_paragraph()

    # ── FITUR / MODUL ──────────────────────────────────
    add_heading_numbered(doc, "Fitur / Modul Website")
    doc.add_paragraph()

    labels = ["a", "b", "c", "d", "e", "f"]
    for idx, grup in enumerate(FITUR_MODUL):
        label_huruf = labels[idx]
        add_heading_numbered(doc, f"{label_huruf}. {grup['label']}", bold=True)
        for fitur, deskripsi in grup["items"]:
            p = doc.add_paragraph(style='List Bullet')
            run_f = p.add_run(f"{fitur} : ")
            set_font(run_f, bold=True, size=12)
            run_d = p.add_run(deskripsi)
            set_font(run_d, bold=False, size=12)
    doc.add_paragraph()

    # ── PERANCANGAN ────────────────────────────────────
    doc.add_page_break()
    add_heading_numbered(doc, "Perancangan", bold=True, size=13)
    doc.add_paragraph()

    for modul in PERANCANGAN_MODUL:
        # Nama modul
        add_heading_numbered(doc, f"Modul: {modul['nama']}", bold=True)
        doc.add_paragraph()

        # A. Identifikasi User
        add_heading_numbered(doc, "Identifikasi User", bold=False)
        table_user = doc.add_table(rows=1, cols=3)
        set_table_style(table_user)
        add_table_row(table_user, ["No", "Aktor", "Deskripsi"], is_header=True)
        for i, (aktor, desk) in enumerate(modul["identifikasi_user"], 1):
            add_table_row(table_user, [str(i), aktor, desk])
        doc.add_paragraph()

        # B. Daftar Kebutuhan / Use Case
        add_heading_numbered(doc, "Daftar Kebutuhan / Use Case", bold=False)
        table_uc = doc.add_table(rows=1, cols=3)
        set_table_style(table_uc)
        add_table_row(table_uc, ["No", "Kebutuhan", "Aktor"], is_header=True)
        for i, (kebutuhan, aktor) in enumerate(modul["use_case"], 1):
            add_table_row(table_uc, [str(i), kebutuhan, aktor])
        doc.add_paragraph()

        # C. Use Case Diagram (placeholder)
        add_heading_numbered(doc, "Use Case Diagram", bold=False)
        add_placeholder_image(doc, f"Diagram Use Case – {modul['nama']} (Terlampir)")
        doc.add_paragraph()

        # D. Flowchart Sistem (placeholder)
        add_heading_numbered(doc, "Flowchart Sistem", bold=False)
        add_placeholder_image(doc, f"Flowchart Sistem – {modul['nama']} (Terlampir)")
        doc.add_paragraph()

        # E. Perancangan Database
        add_heading_numbered(doc, "Perancangan Database", bold=False)

        # Normalisasi
        add_para(doc, "Tahapan Normalisasi Database", bold=False, indent=0.5)
        for nama_tabel, kolom_list in modul["normalisasi"]:
            add_para(doc, f"Tabel: {nama_tabel}", bold=True, indent=1.0)
            table_norm = doc.add_table(rows=1, cols=3)
            set_table_style(table_norm)
            add_table_row(table_norm, ["Atribut", "Tipe Data", "Keterangan"], is_header=True)
            for attr, tipe, ket in kolom_list:
                add_table_row(table_norm, [attr, tipe, ket])
            doc.add_paragraph()

        # Relasi Tabel
        add_para(doc, "Relasi Tabel", bold=False, indent=0.5)
        table_rel = doc.add_table(rows=1, cols=5)
        set_table_style(table_rel)
        add_table_row(table_rel, ["Tabel Asal", "Atribut FK", "Tabel Tujuan", "Atribut PK", "Jenis Relasi"], is_header=True)
        for row in modul["relasi"]:
            add_table_row(table_rel, list(row))
        doc.add_paragraph()
        doc.add_paragraph()

    # ── TAHAPAN PENGEMBANGAN ───────────────────────────
    doc.add_page_break()
    add_heading_numbered(doc, "Tahapan Proses Pengembangan Website", bold=True)
    doc.add_paragraph()

    tahapan = [
        ("1. Analisis Kebutuhan", "Identifikasi masalah: digitalisasi pembayaran SPP dan transparansi penggalangan dana. Pengumpulan data melalui diskusi kelompok dan studi literatur platform sejenis (GoFundMe, Akseleran)."),
        ("2. Perancangan Sistem", "Desain arsitektur multi-tenant (school_id sebagai isolator data), ERD database PostgreSQL via Supabase, wireframe UI menggunakan Figma, dan alur pembayaran Xendit."),
        ("3. Implementasi Frontend", "Pengembangan komponen React 18 + TypeScript menggunakan Vite 6. Styling menggunakan Tailwind CSS v4. Komponen utama: AuthContext, ProtectedRoute, DashboardSiswa, SchoolBillsPage, DonorCampaignsPage."),
        ("4. Implementasi Backend", "Konfigurasi Supabase: skema tabel, RLS (Row Level Security), seed data. Pembuatan Edge Functions: xendit-create-invoice, xendit-webhook, whatsapp-send."),
        ("5. Integrasi & Pengujian", "Integrasi Xendit Payment Gateway (mode sandbox), pengujian alur pembayaran end-to-end, pengujian multi-role login, dan validasi form. Testing menggunakan metode Black-box."),
        ("6. Deployment", "Build production menggunakan Vite. Deploy ke Vercel menggunakan konfigurasi vercel.json dengan SPA rewrite rules. URL production: https://edufin-app.vercel.app"),
    ]

    for judul_t, isi_t in tahapan:
        add_para(doc, judul_t, bold=True)
        add_para(doc, isi_t, bold=False, indent=0.5)
        doc.add_paragraph()

    # ── URL & DOKUMENTASI ──────────────────────────────
    add_heading_numbered(doc, "URL dan Dokumentasi", bold=True)
    doc.add_paragraph()

    urls = [
        ("Aplikasi (Production)", "https://edufin-app.vercel.app"),
        ("Repository GitHub", "https://github.com/stanlevv/edufin-app"),
        ("Dokumentasi API Supabase", "Tersedia di folder /docs/features/ (9 file Markdown)"),
        ("Figma Design", "[Terlampir – URL Figma]"),
    ]

    table_url = doc.add_table(rows=1, cols=2)
    set_table_style(table_url)
    add_table_row(table_url, ["Kategori", "URL / Keterangan"], is_header=True)
    for kat, url in urls:
        add_table_row(table_url, [kat, url])
    doc.add_paragraph()

    # ── PENGUJIAN ──────────────────────────────────────
    add_heading_numbered(doc, "Laporan Pengujian Website dengan Metode Usability", bold=True)
    add_para(doc, "[Terlampir – Hasil kuesioner SUS (System Usability Scale) dari 5–10 responden]", bold=False)
    doc.add_paragraph()

    add_heading_numbered(doc, "Laporan Pengujian Website dengan Metode Black-box Testing", bold=True)

    blackbox_data = [
        ("Login multi-role (siswa, sekolah, donatur)", "Email & password valid per role", "Redirect ke halaman sesuai role masing-masing", "✅ Berhasil"),
        ("Buat tagihan SPP massal (per kelas)", "Pilih kelas 7A, nominal Rp 350.000, due date 15 Juli 2026", "Tagihan muncul di daftar siswa kelas 7A", "✅ Berhasil"),
        ("Upload bukti transfer manual", "Upload foto JPG bukti transfer", "File tersimpan, status tagihan = 'Menunggu Verifikasi'", "✅ Berhasil"),
        ("Pengajuan cicilan 3x", "Pilih tagihan, pilih 3 periode, isi alasan", "Status tagihan berubah menjadi 'Cicilan', 3 sub-tagihan terbuat", "✅ Berhasil"),
        ("Donasi ke kampanye (Xendit sandbox)", "Isi form donasi Rp 50.000, pilih QRIS", "Redirect ke halaman pembayaran Xendit (mode sandbox)", "✅ Berhasil"),
        ("Akses halaman protected tanpa login", "Akses /school/bills tanpa sesi aktif", "Redirect otomatis ke halaman /login", "✅ Berhasil"),
        ("Import siswa massal via CSV", "Upload file CSV berisi 30 data siswa", "Preview tabel muncul, data tersimpan setelah konfirmasi", "✅ Berhasil"),
    ]

    doc.add_paragraph()
    table_bb = doc.add_table(rows=1, cols=4)
    set_table_style(table_bb)
    add_table_row(table_bb, ["Skenario Pengujian", "Input", "Output yang Diharapkan", "Hasil"], is_header=True)
    for row in blackbox_data:
        add_table_row(table_bb, list(row))

    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "— Akhir Laporan —", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── SIMPAN ─────────────────────────────────────────
    output_path = r"C:\Users\diego armando\Documents\a evil twin\File MIT\kkksss\LAPORAN_EDUFIN_UTS.docx"
    doc.save(output_path)
    print("\nLaporan berhasil dibuat!")
    print(f"Lokasi: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_laporan()
