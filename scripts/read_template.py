"""
Script untuk membaca dan menganalisis struktur template Word.
"""
from docx import Document
from docx.shared import Pt, RGBColor
import json

template_path = r"C:\Users\diego armando\Documents\a evil twin\File MIT\kkksss\Format-Laporan-Project-Final_rev.docx"

doc = Document(template_path)

print("=" * 60)
print("ANALISIS TEMPLATE WORD")
print("=" * 60)

# Cek styles yang tersedia
print("\n[STYLES YANG TERSEDIA]")
style_names = [s.name for s in doc.styles]
heading_styles = [s for s in style_names if s and ('heading' in s.lower() or 'Heading' in s)]
print(f"Heading styles: {heading_styles}")

# Baca semua paragraph
print("\n[ISI DOKUMEN - 100 PARAGRAF PERTAMA]")
for i, para in enumerate(doc.paragraphs[:100]):
    if para.text.strip():
        style = para.style.name
        font_size = None
        font_name = None
        bold = None
        try:
            if para.runs:
                run = para.runs[0]
                font_size = run.font.size.pt if run.font.size else None
                font_name = run.font.name
                bold = run.font.bold
        except:
            pass
        print(f"[{i}] Style='{style}' | Bold={bold} | Size={font_size}pt | Font={font_name}")
        print(f"     Text: {para.text[:100]}")
        print()

# Cek sections / page settings
print("\n[PAGE SETTINGS]")
section = doc.sections[0]
print(f"Page width: {section.page_width.cm:.2f} cm")
print(f"Page height: {section.page_height.cm:.2f} cm")
print(f"Left margin: {section.left_margin.cm:.2f} cm")
print(f"Right margin: {section.right_margin.cm:.2f} cm")
print(f"Top margin: {section.top_margin.cm:.2f} cm")
print(f"Bottom margin: {section.bottom_margin.cm:.2f} cm")

# Cek tabel jika ada
print(f"\n[TABEL] Jumlah tabel dalam template: {len(doc.tables)}")
for i, table in enumerate(doc.tables[:3]):
    print(f"\nTabel {i+1}: {len(table.rows)} baris x {len(table.columns)} kolom")
    for j, row in enumerate(table.rows[:3]):
        for k, cell in enumerate(row.cells[:4]):
            print(f"  [{j}][{k}]: {cell.text[:50]}")
